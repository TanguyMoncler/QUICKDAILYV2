from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn

# Couleurs
NAME_COLOR = RGBColor(0x00, 0x70, 0xC0)  # #0070C0
GREEN = RGBColor(0x00, 0xB0, 0x50)       # #00B050
RED = RGBColor(0xC0, 0x00, 0x00)         # #C00000
BLACK = RGBColor(0x00, 0x00, 0x00)

# -------------------------------
# FONCTIONS UTILITAIRES
# -------------------------------

def format_pct(value):
    """Formate le pourcentage et retourne le texte à afficher (2 décimales)"""
    if value is None or abs(value) < 0.000001:
        return "-"
    elif value > 0:
        return f"+{value:.2f}%"
    else:
        return f"{value:.2f}%"


def apply_cell_style(cell_or_paragraph, color=None):
    """Applique la police, taille et couleur à toutes les runs d'une cellule ou d'un paragraphe"""
    if hasattr(cell_or_paragraph, "paragraphs"):
        iterable = cell_or_paragraph.paragraphs
    else:
        iterable = [cell_or_paragraph]

    for paragraph in iterable:
        for run in paragraph.runs:
            run.font.name = "Sitka Display"
            run.font.size = Pt(8)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sitka Display")
            if color:
                run.font.color.rgb = color


def replace_placeholder_in_doc(doc, placeholder, value, color=None):
    """
    Remplace un placeholder {{...}} dans TOUS les tableaux du document
    et applique le style (police, taille, couleur)
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, value)
                    apply_cell_style(cell, color)


def replace_in_doc(doc, placeholder, value, color=None):
    """Remplace placeholder partout dans le document (pas seulement tableaux)"""
    # Paragraphs
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, value)
            apply_cell_style(paragraph, color)

    # Tables
    replace_placeholder_in_doc(doc, placeholder, value, color)


# -------------------------------
# TABLEAU 1 : MOST ACTIVE / BEST / WORST
# -------------------------------
def fill_table_1(doc, market_data):
    most_active = market_data["most_active"]
    best = market_data["best_performers"]
    worst = market_data["worst_performers"]

    for i in range(5):
        ma = most_active[i]
        bp = best[i]
        wp = worst[i]

        # Volume en million x (deux décimales)
        ma_vol = f"{ma['volume'] / 1_000_000:.2f}x"

        # Pourcentages (2 décimales)
        bp_txt = format_pct(bp.get("pct_change"))
        wp_txt = format_pct(wp.get("pct_change"))

        # Couleurs pour pourcentages
        bp_color = GREEN if bp.get("pct_change", 0) > 0 else RED if bp.get("pct_change", 0) < 0 else BLACK
        wp_color = RED if wp.get("pct_change", 0) < 0 else GREEN if wp.get("pct_change", 0) > 0 else BLACK

        # Remplacement placeholders selon le format demandé par l'utilisateur
        replace_placeholder_in_doc(
            doc,
            f"{{{{MOST ACTIVE STOCK {i+1}}}}}",
            ma.get("name", ma.get("ticker", "")).strip(),
            NAME_COLOR
        )
        replace_placeholder_in_doc(
            doc,
            f"{{{{MULTIPLE {i+1}}}}}",
            ma_vol,
            NAME_COLOR
        )
        # Compatibilité avec anciens placeholders
        replace_placeholder_in_doc(
            doc,
            f"{{{{MAS MULTIPLE {i+1}}}}}",
            ma_vol,
            NAME_COLOR
        )
        replace_placeholder_in_doc(
            doc,
            f"{{{{BEST PERFORMER {i+1}}}}}",
            bp.get("name", bp.get("ticker", "")).strip(),
            NAME_COLOR
        )
        replace_placeholder_in_doc(
            doc,
            f"{{{{INCREASE {i+1}}}}}",
            bp_txt,
            bp_color
        )
        replace_placeholder_in_doc(
            doc,
            f"{{{{WORST PERFORMER {i+1}}}}}",
            wp.get("name", wp.get("ticker", "")).strip(),
            NAME_COLOR
        )
        replace_placeholder_in_doc(
            doc,
            f"{{{{DECREASE {i+1}}}}}",
            wp_txt,
            wp_color
        )


# -------------------------------
# TABLEAU 2 : INDICES / FX / COMMODITIES
# -------------------------------
def fill_table_2(doc, table_2_data):
    for item in table_2_data:
        placeholder = item["placeholder"]
        display = item["display"]
        pct = item.get("pct_change", 0)
        color = GREEN if pct > 0 else RED if pct < 0 else BLACK
        replace_placeholder_in_doc(doc, placeholder, display, None)

    # Remplacer les MvtX (Mvt1..Mvt9)
    # On suppose que les placeholders {{Mvt1}}, {{Mvt2}} ... sont dans l'ordre fourni
    for idx, item in enumerate(table_2_data, start=1):
        mvt_placeholder = f"{{{{Mvt{idx}}}}}"
        mvt_txt = format_pct(item.get("pct_change"))
        mvt_color = GREEN if item.get("pct_change", 0) > 0 else RED if item.get("pct_change", 0) < 0 else BLACK
        replace_placeholder_in_doc(doc, mvt_placeholder, mvt_txt, mvt_color)


# -------------------------------
# TABLEAU 3 : SECTEURS EUROSTOXX
# -------------------------------
def fill_table_3(doc, table_3_data):
    # Remplacer secteurs (closing price) and Mvt10..Mvt14
    for item in table_3_data:
        replace_placeholder_in_doc(doc, item["placeholder"], item["display"], None)

    for idx, item in enumerate(table_3_data, start=10):
        mvt_placeholder = f"{{{{Mvt{idx}}}}}"
        mvt_txt = format_pct(item.get("pct_change"))
        mvt_color = GREEN if item.get("pct_change", 0) > 0 else RED if item.get("pct_change", 0) < 0 else BLACK
        replace_placeholder_in_doc(doc, mvt_placeholder, mvt_txt, mvt_color)


# -------------------------------
# DATE
# -------------------------------
def _ordinal(n):
    if 11 <= (n % 100) <= 13:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")


def replace_date(doc, date_obj):
    """Remplace tous les {{DATE}} par une string type "January 12th 2026""" 
    date_str = date_obj.strftime("%B %d %Y")
    # remplacer day with ordinal
    day = date_obj.day
    day_with_suffix = f"{day}{_ordinal(day)}"
    date_str = date_obj.strftime("%B") + f" {day_with_suffix} " + date_obj.strftime("%Y")
    replace_in_doc(doc, "{{DATE}}", date_str, None)


def clear_remaining_placeholders(doc):
    """Remplace tout placeholder {{...}} restant par '-' et applique le style par défaut"""
    import re
    pattern = re.compile(r"\{\{[^}]+\}\}")

    # Paragraphs
    for paragraph in doc.paragraphs:
        if pattern.search(paragraph.text):
            paragraph.text = pattern.sub("-", paragraph.text)
            apply_cell_style(paragraph, BLACK)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if pattern.search(cell.text):
                    cell.text = pattern.sub("-", cell.text)
                    apply_cell_style(cell, BLACK)

